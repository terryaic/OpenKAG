from docx import Document
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 读取 HTML 文件内容
with open('output.html', 'r', encoding='utf-8') as f:
    html_content = f.read()

# 使用 BeautifulSoup 解析 HTML
soup = BeautifulSoup(html_content, 'html.parser')

# 创建一个新的 Word 文档
doc = Document()

# 解析 HTML 表格并将其添加到 Word 文档中
for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'table']):
    if element.name == 'p':
        doc.add_paragraph(element.text)
    elif element.name == 'h1':
        doc.add_heading(element.text, level=1)
    elif element.name == 'h2':
        doc.add_heading(element.text, level=2)
    elif element.name == 'h3':
        doc.add_heading(element.text, level=3)
    elif element.name == 'table':
        # 创建表格并填充内容
        rows = element.find_all('tr')
        cols = len(rows[0].find_all(['td', 'th']))
        table = doc.add_table(rows=1, cols=cols)
        hdr_cells = table.rows[0].cells
        for i, th in enumerate(rows[0].find_all(['td', 'th'])):
            hdr_cells[i].text = th.text

        for row in rows[1:]:
            row_cells = table.add_row().cells
            for i, td in enumerate(row.find_all('td')):
                row_cells[i].text = td.text

        # 设置表格边框样式
        tbl = table._tbl  # 获取底层表格对象
        tblPr = tbl.tblPr  # 获取或创建表格属性

        # 创建并附加边框元素
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 设置边框宽度，单位为1/8磅
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 设置边框颜色为黑色
            tblBorders.append(border)

        # 将边框附加到表格属性
        tblPr.append(tblBorders)

# 保存 Word 文档
doc.save('final_output.docx')

