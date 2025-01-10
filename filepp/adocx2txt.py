import sys
import docx
from docx import Document
from filepp.textutil import content_to_txt

def docx_to_text(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def x_read_docx_in_order(file_path):
    # 打开DOCX文件
    doc = Document(file_path)
    
    # 初始化一个列表来存储所有内容
    content = []
    
    # 遍历文档中的所有元素
    for element in doc.element.body:
        if element.tag.endswith('p'):  # 段落
            para = element
            content.append(('paragraph', para.text))
        elif element.tag.endswith('tbl'):  # 表格
            table = element
            table_data = []
            for row in table.findall('.//w:tr'):
                row_data = []
                for cell in row.findall('.//w:tc'):
                    cell_text = ''.join(node.text for node in cell.findall('.//w:t'))
                    row_data.append(cell_text)
                table_data.append(row_data)
            content.append(('table', table_data))
    
    return content

def read_docx_in_order(file_path):
    # 打开DOCX文件
    doc = Document(file_path)
    
    # 初始化一个列表来存储所有内容
    content = []
    
    # 遍历文档中的所有元素
    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            content.append(('paragraph', block.text))
        elif isinstance(block, docx.table.Table):
            table_data = []
            for row in block.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content.append(('table', table_data))
    
    return content

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

def docx_to_txt(file_path):
  print("doc to txt:"+file_path)
  content = read_docx_in_order(file_path)
  out = content_to_txt(content)
  return out

if __name__ == "__main__":
  # 使用示例
  file_path = sys.argv[1]
  out = docx_to_txt(file_path)
  print(out)
