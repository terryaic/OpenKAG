import sys
import docx
from docx import Document
from filepp.textutil import content_to_txt

def docx_to_text(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def x_read_docx_in_order(file_path):
    # 打开DOCX文件
    doc = Document(file_path)
    
    # 初始化一个列表来存储所有内容
    content = []
    
    # 遍历文档中的所有元素
    for element in doc.element.body:
        if element.tag.endswith('p'):  # 段落
            para = element
            content.append(('paragraph', para.text))
        elif element.tag.endswith('tbl'):  # 表格
            table = element
            table_data = []
            for row in table.findall('.//w:tr'):
                row_data = []
                for cell in row.findall('.//w:tc'):
                    cell_text = ''.join(node.text for node in cell.findall('.//w:t'))
                    row_data.append(cell_text)
                table_data.append(row_data)
            content.append(('table', table_data))
    
    return content

def fill_with_llm(last_text, context):
    import requests
    import json
    prompt = "according the text below to fill the answer.\n"
    prompt += "text:"+ context +"\n"
    prompt += "question:" + last_text + "\n"
    body = {"question": prompt}
    url = "http://192.168.1.13:8993/generate"
    response = requests.post(url, json=body)
    ret = json.loads(response.content.decode("utf8"))
    return ret['answer']

def read_docx_in_order(file_path, out_file_path, context=""):
    # 打开DOCX文件
    doc = Document(file_path)
    
    # 初始化一个列表来存储所有内容
    content = []
    
    # 遍历文档中的所有元素
    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            content.append(('paragraph', block.text))
        elif isinstance(block, docx.table.Table):
            table_data = []
            for row in block.rows:
                for cell in row.cells:
                    if len(cell.text) == 0:
                        text = fill_with_llm(last_text, context)
                        print(text)
                        cell.text= text
                    else:
                        last_text = cell.text

    doc.save(out_file_path)
    
    return content

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

def docx_to_txt(file_path,out_file_path,context_path):
  from fileprocesspipeline import FileProcessPipeline
  pipeline = FileProcessPipeline(None)
  context = pipeline.process_text(context_path)
  if file_path.endswith(".xlsx"):
     sheet_op(file_path,out_file_path, context)
     out = ''
  else:
     content = read_docx_in_order(file_path,out_file_path, context)
     out = content_to_txt(content)
  return out

from openpyxl import load_workbook

def sheet_op(filepath,outpath,context=''):
    # 加载已有的 Excel 工作簿
    workbook = load_workbook(filepath)

    # 选择工作表
    sheet = workbook.active  # 或者 workbook['Sheet1']

    # 读取单元格的值
    #value = sheet['A1'].value
    #print(value)

    # 遍历所有行和列
    for row in sheet.iter_rows():#values_only=True):
        for cell in row:
            if cell.value:
                last_text = cell.value
                print(cell.value)
            else:
                merged=False
                for merged_range in sheet.merged_cells.ranges:
                    if cell.coordinate in merged_range:
                        print("merged cell")
                        merged = True
                        break
                if not merged:
                    cell.value = fill_with_llm(last_text,context)
                    print("fill with text:"+cell.value)
                 
    workbook.save(outpath)

if __name__ == "__main__":
  # 使用示例
  file_path = sys.argv[1]
  context_path = sys.argv[2]
  out = docx_to_txt(file_path, file_path.replace(".","_out."), context_path)
